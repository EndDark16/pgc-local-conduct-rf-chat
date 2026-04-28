"""Genera documentos Word (.docx) desde los contenidos Markdown del proyecto."""

from __future__ import annotations

from pathlib import Path
from typing import Iterable


DOCS_DIR = Path(__file__).resolve().parent
DOC_MD = DOCS_DIR / "documento_pgc_contenido.md"
PRES_MD = DOCS_DIR / "presentacion_pgc_contenido.md"
DOC_OUT = DOCS_DIR / "documento_pgc_generado.docx"
PRES_OUT = DOCS_DIR / "presentacion_pgc_guion_generado.docx"


def _load_docx():
    try:
        from docx import Document  # type: ignore
    except Exception:  # noqa: BLE001
        print("No se pudo importar python-docx. Instala con:")
        print("pip install python-docx")
        return None
    return Document


def _write_md_to_docx(md_path: Path, out_path: Path, Document) -> None:
    text = md_path.read_text(encoding="utf-8")
    doc = Document()
    for raw_line in text.splitlines():
        line = raw_line.rstrip()
        if not line:
            doc.add_paragraph("")
            continue
        if line.startswith("# "):
            doc.add_heading(line[2:].strip(), level=1)
            continue
        if line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=2)
            continue
        if line.startswith("### "):
            doc.add_heading(line[4:].strip(), level=3)
            continue
        if line.startswith("- "):
            doc.add_paragraph(line[2:].strip(), style="List Bullet")
            continue
        if line[:3].isdigit() and line[1:3] == ". ":
            doc.add_paragraph(line[3:].strip(), style="List Number")
            continue
        doc.add_paragraph(line)
    doc.save(out_path)


def main() -> int:
    Document = _load_docx()
    if Document is None:
        return 1

    if not DOC_MD.exists() or not PRES_MD.exists():
        print("Faltan archivos markdown base en docs/.")
        return 1

    _write_md_to_docx(DOC_MD, DOC_OUT, Document)
    _write_md_to_docx(PRES_MD, PRES_OUT, Document)
    print("Documento generado:", DOC_OUT)
    print("Guion de presentación generado:", PRES_OUT)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
